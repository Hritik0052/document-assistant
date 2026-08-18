from pathlib import Path

from asgiref.sync import sync_to_async
from docx import Document as WordDocument

from rag.loaders.base import BaseLoader
from rag.schemas import ExtractedDocument


class DocxLoader(BaseLoader):
    suffix = '.docx'

    async def load(self, path: Path) -> ExtractedDocument:
        return await sync_to_async(self._load_sync)(path)

    def _load_sync(self, path: Path) -> ExtractedDocument:
        document = WordDocument(str(path))
        paragraphs = [para.text for para in document.paragraphs if para.text.strip()]
        text = '\n'.join(paragraphs)
        return ExtractedDocument(
            text=text,
            source=path.name,
            extra={'loader': 'docx', 'paragraphs': len(paragraphs)},
        )
